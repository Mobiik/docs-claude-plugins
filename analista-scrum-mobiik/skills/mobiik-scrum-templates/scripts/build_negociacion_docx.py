"""
Guía de Preparación de Negociación — Mobiik
Documento de preparación previo a una negociación comercial / resolución de conflicto.
Estructura basada en negociación por intereses (Harvard), BATNA y plan de concesiones.
Estándar visual Mobiik: encabezado branded, acento verde lima #AADC1E.

Basado en la plantilla corporativa Mobiik_Template_Estrategias_Negociacion.docx
Dependencia:  pip3 install --user python-docx
Uso:          python3 build_negociacion_docx.py

NOTA: Este script genera el ENTREGABLE de preparación. El análisis de estrategias
(replanteo por intereses, opciones diferenciadas, recomendación) lo produce el
agente en texto; este documento es el soporte formal para preparar la mesa.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── DATOS DE LA NEGOCIACIÓN (editar) ────────────────────────────────
CONTRAPARTE   = "[Cliente / proveedor]"
OBJETO        = "[Servicio / proyecto / contrato]"
VALOR         = "[$ MXN]"
INTERLOCUTORES = "[Quién decide / quién influye]"

INTERESES_MOBIIK = [
    "[Qué queremos realmente lograr — más allá del precio: relación, recurrencia, referencias, margen]",
]
INTERESES_CONTRA = [
    "[Qué necesita / qué le preocupa: presupuesto, tiempos, riesgo, percepción interna]",
]

POSICION_IDEAL  = "[Mejor escenario]"
OBJETIVO_REAL   = "[Lo esperable]"
LINEA_ROJA      = "[Límite — abajo de esto, no hay trato]"
BATNA_MOBIIK    = "[____]"
BATNA_CONTRA    = "[____]"

# (Variable, Costo para Mobiik, Valor para cliente)
VARIABLES = [
    ("Precio / forma de pago", "[Bajo/Medio/Alto]", "[Bajo/Medio/Alto]"),
    ("Plazos / tiempos de entrega", "[ ]", "[ ]"),
    ("Alcance / entregables", "[ ]", "[ ]"),
    ("Garantía / soporte", "[ ]", "[ ]"),
    ("Exclusividad / recurrencia", "[ ]", "[ ]"),
]

# (Objeción probable, Respuesta preparada)
OBJECIONES = [
    ("“Está caro”", "[Reencuadrar a valor / TCO / riesgo evitado]"),
    ("“Otro proveedor ofrece menos”", "[Diferenciadores — experiencia, soporte, track record]"),
]

# (Orden, Concesión, Contrapartida solicitada)
CONCESIONES = [
    ("1", "[Concesión]", "[Contrapartida]"),
    ("2", "[Concesión]", "[Contrapartida]"),
    ("3", "[Concesión]", "[Contrapartida]"),
]

SENALES_CIERRE   = "[____]"
COMPROMISOS_DOC  = "[____]"
RESP_FORMALIZAR  = "[____]"

OUT = "./entregables/Preparacion Negociacion - <contraparte>.docx"
# ─────────────────────────────────────────────────────────────────────

# ─── Brand ───────────────────────────────────────────────────────────
LIME      = RGBColor(0xAA, 0xDC, 0x1E)
BLACK     = RGBColor(0x0A, 0x0A, 0x0A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY      = RGBColor(0x66, 0x66, 0x66)
FONT      = "Arial"

doc = Document()
sec = doc.sections[0]
sec.page_height = Inches(11); sec.page_width = Inches(8.5)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)
style = doc.styles['Normal']; style.font.name = FONT; style.font.size = Pt(11)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), hex_color); tc_pr.append(shd)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), size); b.set(qn('w:color'), color)
        tcb.append(b)
    tc_pr.append(tcb)

def add_heading_bar(text, num=None, *, size=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if num:
        rn = p.add_run(f"{num}  ")
        rn.font.name = FONT; rn.font.size = Pt(size); rn.font.bold = True; rn.font.color.rgb = LIME
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(size); run.font.bold = True; run.font.color.rgb = BLACK
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'AADC1E')
    pBdr.append(left); p_pr.append(pBdr)
    return p

def add_paragraph(text, *, italic=False, color=None, size=11, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size); r.font.italic = italic
    if color: r.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(11)
    return p

def add_kv_grid(rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    for ri, (k, v) in enumerate(rows):
        kc, vc = table.rows[ri].cells
        kc.width = Inches(2.3); vc.width = Inches(4.2)
        set_cell_bg(kc, "0A0A0A"); set_cell_border(kc)
        set_cell_bg(vc, "F5F5F5" if ri % 2 == 0 else "FFFFFF"); set_cell_border(vc)
        kc.vertical_alignment = vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        kc.text = ""; vc.text = ""
        rk = kc.paragraphs[0].add_run(k)
        rk.font.name = FONT; rk.font.size = Pt(10); rk.font.bold = True; rk.font.color.rgb = WHITE
        rv = vc.paragraphs[0].add_run(v); rv.font.name = FONT; rv.font.size = Pt(10.5)
    return table

def add_branded_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr[i]; set_cell_bg(c, "0A0A0A"); set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.name = FONT; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            c = cells[ci]; set_cell_bg(c, "F5F5F5" if ri % 2 == 0 else "FFFFFF"); set_cell_border(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; c.text = ""
            r = c.paragraphs[0].add_run(str(val)); r.font.name = FONT; r.font.size = Pt(10)
    return table


# ═══════════════════════════════════════════════════════════════════
# ENCABEZADO
# ═══════════════════════════════════════════════════════════════════
title = doc.add_paragraph(); title.paragraph_format.space_after = Pt(2)
rt = title.add_run("Estrategias de Negociación")
rt.font.name = FONT; rt.font.size = Pt(24); rt.font.bold = True; rt.font.color.rgb = BLACK
p_pr = title._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24'); left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'AADC1E')
pBdr.append(left); p_pr.append(pBdr)

sub = doc.add_paragraph(); sub.paragraph_format.space_after = Pt(16)
rs = sub.add_run("Guía de preparación para negociaciones comerciales — Mobiik")
rs.font.name = FONT; rs.font.size = Pt(11); rs.font.italic = True; rs.font.color.rgb = GREY

add_paragraph(
    "Este documento sirve para preparar una negociación antes de sentarse a la mesa. "
    "Complete cada sección con la información del trato específico; las preguntas guía "
    "ayudan a no dejar cabos sueltos.", after=12)

# 1. Contexto del trato
add_heading_bar("Contexto del trato", num="1")
add_kv_grid([
    ("Contraparte", CONTRAPARTE),
    ("Objeto de la negociación", OBJETO),
    ("Valor estimado", VALOR),
    ("Interlocutores clave", INTERLOCUTORES),
])

# 2. Objetivos e intereses
add_heading_bar("Objetivos e intereses", num="2")
add_paragraph("Nuestros intereses (Mobiik)", color=BLACK, after=2)
for it in INTERESES_MOBIIK:
    add_bullet(it)
add_paragraph("Intereses de la contraparte", color=BLACK, after=2)
for it in INTERESES_CONTRA:
    add_bullet(it)

# 3. Márgenes y alternativas
add_heading_bar("Márgenes y alternativas", num="3")
add_branded_table(
    ["Posición ideal", "Objetivo realista", "Línea roja (mínimo)"],
    [(POSICION_IDEAL, OBJETIVO_REAL, LINEA_ROJA)],
    col_widths=[Inches(2.2), Inches(2.2), Inches(2.1)],
)
add_paragraph(f"Nuestra BATNA (mejor alternativa si no hay acuerdo): {BATNA_MOBIIK}", after=2)
add_paragraph(f"BATNA estimada de la contraparte: {BATNA_CONTRA}", after=8)

# 4. Variables de intercambio
add_heading_bar("Variables de intercambio", num="4")
add_paragraph(
    "Identifique qué puede mover además del precio. Conceder en variables de bajo "
    "costo para Mobiik que tengan alto valor para el cliente.", after=6)
add_branded_table(
    ["Variable", "Costo para Mobiik", "Valor para cliente"],
    VARIABLES,
    col_widths=[Inches(2.8), Inches(1.8), Inches(1.9)],
)

# 5. Tácticas recomendadas
add_heading_bar("Tácticas recomendadas", num="5")
add_paragraph("Antes de la mesa", color=BLACK, after=2)
add_bullet("Ancla con una propuesta de valor, no con el precio: abre justificando el porqué antes del cuánto.")
add_bullet("Prepara 2–3 paquetes (bueno / mejor / óptimo) para dar opción sin bajar el piso.")
add_paragraph("Durante", color=BLACK, after=2)
add_bullet("Escucha primero: deja que la contraparte revele restricciones y prioridades.")
add_bullet("Condiciona toda concesión (“si… entonces…”); nunca regales sin contrapartida.")
add_bullet("Mantén el tono colaborativo: ataca el problema, no a la persona.")
add_paragraph("Manejo de objeciones", color=BLACK, after=2)
add_branded_table(
    ["Objeción probable", "Respuesta preparada"],
    OBJECIONES,
    col_widths=[Inches(2.6), Inches(3.9)],
)

# 6. Plan de concesiones
add_heading_bar("Plan de concesiones", num="6")
add_paragraph(
    "Define el orden y el límite de lo que estás dispuesto a ceder, y qué pides a "
    "cambio en cada paso.", after=6)
add_branded_table(
    ["Orden", "Concesión", "Contrapartida solicitada"],
    CONCESIONES,
    col_widths=[Inches(0.9), Inches(2.8), Inches(2.8)],
)

# 7. Cierre y siguientes pasos
add_heading_bar("Cierre y siguientes pasos", num="7")
add_bullet(f"Señales de cierre a buscar: {SENALES_CIERRE}")
add_bullet(f"Compromisos a documentar de inmediato: {COMPROMISOS_DOC}")
add_bullet(f"Responsable de formalizar (minuta / contrato): {RESP_FORMALIZAR}")

# Pie
doc.add_paragraph()
pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run("© 2026 Mobiik   |   coding for the future   |   Documento Confidencial")
rf.font.name = FONT; rf.font.size = Pt(9); rf.font.italic = True; rf.font.color.rgb = GREY


# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
import os
os.makedirs(os.path.dirname(OUT) or ".", exist_ok=True)
doc.save(OUT)
print(f"✓ Saved: {OUT}")
