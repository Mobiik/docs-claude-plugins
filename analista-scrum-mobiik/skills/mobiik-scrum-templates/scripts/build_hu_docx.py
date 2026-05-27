"""
Documento de Historias de Usuario — Migración Preselección TecMilenio
Formal entregable al cliente con sign-off.
Estándar visual Mobiik: portada oscura, headers con barra lateral verde, tablas branded.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand
LIME = RGBColor(0xAA, 0xDC, 0x1E)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)
DARK_GREY = RGBColor(0x25, 0x25, 0x25)

FONT = "Arial"

doc = Document()

# Page setup: Letter, 1" margins
sec = doc.sections[0]
sec.page_height = Inches(11)
sec.page_width = Inches(8.5)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

# Default font
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(11)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_heading_bar(doc, text, *, size=18, color=BLACK, bar_color=LIME):
    """H1 with a green bar accent to its left (no underline)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    # Left border on paragraph (green bar)
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'AADC1E')
    pBdr.append(left)
    p_pr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = DARK_GREY
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = BLACK
    return p

def add_paragraph(doc, text, *, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = FONT
        r1.font.size = Pt(10.5)
        r1.font.bold = True
        r2 = p.add_run(text)
        r2.font.name = FONT
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(10.5)
    return p

def add_branded_table(doc, headers, rows, *, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        set_cell_bg(cell, "0A0A0A")
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
    # Rows
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cell = row_cells[ci]
            set_cell_bg(cell, "F5F5F5" if ri % 2 == 0 else "FFFFFF")
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = FONT
            r.font.size = Pt(10)
    return table


# ═══ PORTADA ═══════════════════════════════════════════════════════
# Dark "block" simulated with a single-cell table
cover_table = doc.add_table(rows=1, cols=1)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_cell = cover_table.rows[0].cells[0]
cover_cell.width = Inches(6.5)
set_cell_bg(cover_cell, "0A0A0A")

# Add content inside cover cell
cover_cell.text = ""
# spacer
p_sp = cover_cell.paragraphs[0]
p_sp.add_run("\n\n\n\n").font.size = Pt(12)

# "DOCUMENTO" eyebrow
p1 = cover_cell.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = p1.add_run("DOCUMENTO TÉCNICO ENTREGABLE")
r1.font.name = FONT
r1.font.size = Pt(11)
r1.font.bold = True
r1.font.color.rgb = LIME

# Big title
p2 = cover_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r2 = p2.add_run("\nHISTORIAS DE")
r2.font.name = FONT
r2.font.size = Pt(34)
r2.font.bold = True
r2.font.color.rgb = WHITE

p3 = cover_cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
r3 = p3.add_run("USUARIO")
r3.font.name = FONT
r3.font.size = Pt(34)
r3.font.bold = True
r3.font.color.rgb = LIME

# Project
p4 = cover_cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
r4 = p4.add_run("\nMigración Técnica de Preselección")
r4.font.name = FONT
r4.font.size = Pt(15)
r4.font.color.rgb = WHITE

p5 = cover_cell.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
r5 = p5.add_run("Lift-and-Shift a Microsoft Azure PaaS")
r5.font.name = FONT
r5.font.size = Pt(12)
r5.font.italic = True
r5.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# spacer
p_sp2 = cover_cell.add_paragraph()
p_sp2.add_run("\n\n\n").font.size = Pt(12)

# Client / Provider
p6 = cover_cell.add_paragraph()
r6a = p6.add_run("CLIENTE:  ")
r6a.font.name = FONT; r6a.font.size = Pt(10); r6a.font.bold = True; r6a.font.color.rgb = LIME
r6b = p6.add_run("Tecmilenio")
r6b.font.name = FONT; r6b.font.size = Pt(10); r6b.font.color.rgb = WHITE

p7 = cover_cell.add_paragraph()
r7a = p7.add_run("PROVEEDOR:  ")
r7a.font.name = FONT; r7a.font.size = Pt(10); r7a.font.bold = True; r7a.font.color.rgb = LIME
r7b = p7.add_run("Mobiik")
r7b.font.name = FONT; r7b.font.size = Pt(10); r7b.font.color.rgb = WHITE

p8 = cover_cell.add_paragraph()
r8a = p8.add_run("FECHA:  ")
r8a.font.name = FONT; r8a.font.size = Pt(10); r8a.font.bold = True; r8a.font.color.rgb = LIME
r8b = p8.add_run("Mayo 2026")
r8b.font.name = FONT; r8b.font.size = Pt(10); r8b.font.color.rgb = WHITE

p9 = cover_cell.add_paragraph()
r9a = p9.add_run("VERSIÓN:  ")
r9a.font.name = FONT; r9a.font.size = Pt(10); r9a.font.bold = True; r9a.font.color.rgb = LIME
r9b = p9.add_run("v1.0")
r9b.font.name = FONT; r9b.font.size = Pt(10); r9b.font.color.rgb = WHITE

# spacer
p_sp3 = cover_cell.add_paragraph()
p_sp3.add_run("\n\n").font.size = Pt(12)

# Bottom: confidencial + slogan
p10 = cover_cell.add_paragraph()
p10.alignment = WD_ALIGN_PARAGRAPH.LEFT
r10 = p10.add_run("CONFIDENCIAL — Documento sujeto a sign-off")
r10.font.name = FONT; r10.font.size = Pt(9); r10.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA); r10.font.italic = True

p11 = cover_cell.add_paragraph()
p11.alignment = WD_ALIGN_PARAGRAPH.LEFT
r11 = p11.add_run("© 2026 Mobiik   |   coding for the future")
r11.font.name = FONT; r11.font.size = Pt(9); r11.font.color.rgb = LIME

# Page break after cover
doc.add_page_break()


# ═══ CONTROL DE VERSIONES ══════════════════════════════════════════
add_heading_bar(doc, "Control de Versiones", size=18)
add_branded_table(doc,
    ["Versión", "Fecha", "Autor", "Cambios", "Estado", "Aprobado por"],
    [["v1.0", "25-may-2026", "Miroslava Jiménez", "Versión inicial del documento", "Borrador", "PO TecMilenio"]],
    col_widths=[Inches(0.7), Inches(1.0), Inches(1.4), Inches(1.6), Inches(0.8), Inches(1.3)]
)


# ═══ ÍNDICE ════════════════════════════════════════════════════════
add_heading_bar(doc, "Contenido", size=18)
toc_items = [
    "1. Introducción y propósito",
    "2. Audiencia",
    "3. Glosario",
    "4. Actores del sistema",
    "5. Mapa de Épicas",
    "6. Historias de Usuario por Épica",
    "7. Matriz de Trazabilidad Comercial ↔ Funcional",
    "8. Supuestos y Restricciones",
    "9. Fuera de Alcance",
    "10. Anexos",
    "11. Aceptación y Sign-off",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.name = FONT
    p.runs[0].font.size = Pt(11)

doc.add_page_break()


# ═══ 1. INTRODUCCIÓN ═══════════════════════════════════════════════
add_heading_bar(doc, "1. Introducción y propósito", size=18)
add_paragraph(doc,
    "Este documento formaliza las Historias de Usuario (HU) que componen el alcance funcional del proyecto "
    "Migración Técnica del aplicativo Preselección de Tecmilenio hacia infraestructura Azure PaaS bajo el "
    "modelo Lift-and-Shift, ejecutado por Mobiik según la propuesta comercial firmada con fecha 20 de mayo de 2026."
)
add_paragraph(doc,
    "El propósito del documento es:")
add_bullet(doc, "Establecer un entendimiento común y firmado entre Tecmilenio y Mobiik sobre el alcance funcional", bold_prefix="• ")
add_bullet(doc, "Documentar criterios de aceptación claros (Gherkin) por cada HU", bold_prefix="• ")
add_bullet(doc, "Mantener trazabilidad entre los requerimientos comerciales (propuesta), las HU y las tareas en Jira (TEC26)", bold_prefix="• ")
add_bullet(doc, "Servir de referencia durante UAT y como base del sign-off de aceptación", bold_prefix="• ")


# ═══ 2. AUDIENCIA ══════════════════════════════════════════════════
add_heading_bar(doc, "2. Audiencia", size=18)
add_branded_table(doc,
    ["Rol", "Responsabilidad sobre este documento"],
    [
        ["Product Owner TecMilenio",     "Aprobar las HU y firmar sign-off final"],
        ["Project Manager TecMilenio",   "Coordinar UAT y validación contractual"],
        ["Encargado Infraestructura TM", "Validar viabilidad técnica y dependencias"],
        ["Scrum Master Mobiik",          "Mantener trazabilidad HU ↔ Jira ↔ entregables"],
        ["Equipo Técnico Mobiik",         "Implementar y validar las HU"],
    ],
    col_widths=[Inches(2.2), Inches(4.3)]
)


# ═══ 3. GLOSARIO ═══════════════════════════════════════════════════
add_heading_bar(doc, "3. Glosario", size=18)
add_branded_table(doc,
    ["Término", "Definición"],
    [
        ["Lift-and-Shift",        "Estrategia de migración que traslada un aplicativo desde su entorno actual a uno nuevo sin modificar la lógica de negocio"],
        ["PaaS",                   "Platform as a Service. Servicios administrados por el proveedor cloud (App Service, Azure SQL, etc.)"],
        ["IaaS",                   "Infrastructure as a Service. Máquinas virtuales y red administradas por el cliente"],
        ["HU",                     "Historia de Usuario — descripción narrativa de funcionalidad desde la perspectiva del usuario"],
        ["DoD",                    "Definition of Done — criterios objetivos que indican que una HU está terminada"],
        ["Gherkin",                "Lenguaje estructurado para criterios de aceptación (Dado/Cuando/Entonces)"],
        ["MoSCoW",                 "Priorización: Must / Should / Could / Won't have"],
        ["INVEST",                 "Criterios de calidad de HU: Independent, Negotiable, Valuable, Estimable, Small, Testable"],
        ["UAT",                    "User Acceptance Testing — pruebas de aceptación ejecutadas por el cliente"],
        ["RAID",                   "Risks, Assumptions, Issues, Dependencies"],
        ["Smoke Test",             "Verificación rápida de operación básica del aplicativo post-cambio"],
        ["Reescaneo",              "Re-ejecución de análisis de seguridad para validar mitigación de hallazgos"],
        ["TEC26",                  "Proyecto Jira que aloja el backlog completo de esta migración"],
        ["Sprint",                 "Ciclo iterativo Scrum (semanal en este proyecto)"],
        ["Hito de Facturación",   "Punto de control comercial que libera un pago contractual"],
    ],
    col_widths=[Inches(1.8), Inches(4.7)]
)


# ═══ 4. ACTORES ════════════════════════════════════════════════════
add_heading_bar(doc, "4. Actores del sistema", size=18)
add_paragraph(doc,
    "Los siguientes roles intervienen en la ejecución del proyecto. Cada Historia de Usuario referencia el actor que la realiza."
)
add_branded_table(doc,
    ["Actor", "Organización", "Responsabilidad"],
    [
        ["Scrum Master",                   "Mobiik",     "Coordinación, RAID, comunicación con PO"],
        ["Arquitecto de Soluciones",       "Mobiik",     "Análisis técnico, arquitectura objetivo"],
        ["Desarrollador Fullstack",        "Mobiik",     "Migración del aplicativo, configuraciones"],
        ["DevOps / Infraestructura Azure", "Mobiik",     "Provisión Azure, pipelines, conectividad"],
        ["QA Tester",                       "Mobiik",     "Smoke tests, validación funcional, UAT support"],
        ["SecOps",                         "Mobiik",     "Atención de hallazgos de seguridad, reescaneo"],
        ["Product Owner",                  "TecMilenio", "Decisiones de alcance, priorización, aceptación"],
        ["Project Manager",                "TecMilenio", "Coordinación contractual, hitos de facturación"],
        ["Encargado de Infraestructura",   "TecMilenio", "Accesos Azure, GitHub, infraestructura productiva"],
        ["Equipo de Pruebas",               "TecMilenio", "Ejecución de UAT (3 días en Sprint 3)"],
    ],
    col_widths=[Inches(1.8), Inches(1.3), Inches(3.4)]
)


# ═══ 5. MAPA DE ÉPICAS ═════════════════════════════════════════════
add_heading_bar(doc, "5. Mapa de Épicas", size=18)
add_paragraph(doc,
    "El alcance funcional se estructura en 8 épicas. Cada épica agrupa Historias de Usuario relacionadas y "
    "es trazable a una sección de la propuesta comercial."
)
add_branded_table(doc,
    ["#", "Épica", "Jira", "Sprint", "# HU"],
    [
        ["1", "Análisis Técnico y Planeación",         "TEC26-1", "S1",      "4"],
        ["2", "Preparación de Entornos Azure",          "TEC26-2", "S1-S2",   "6"],
        ["3", "Migración Lift-and-Shift del Aplicativo","TEC26-3", "S2",      "5"],
        ["4", "Migración de Bases de Datos MS SQL",     "TEC26-4", "S2",      "4"],
        ["5", "Pipelines CI/CD",                         "TEC26-5", "S2-S3",   "3"],
        ["6", "Atención de 13 Hallazgos de Seguridad",   "TEC26-6", "S3",      "3 + 13 subt"],
        ["7", "Pruebas y UAT",                           "TEC26-7", "S3",      "4"],
        ["8", "Go-Live y Cierre del Proyecto",           "TEC26-8", "Cierre",  "3"],
    ],
    col_widths=[Inches(0.4), Inches(2.8), Inches(0.9), Inches(0.9), Inches(1.5)]
)

doc.add_page_break()


# ═══ 6. HISTORIAS DE USUARIO ═══════════════════════════════════════
add_heading_bar(doc, "6. Historias de Usuario", size=18)
add_paragraph(doc,
    "A continuación se presentan las 32 Historias de Usuario que componen el alcance funcional. "
    "Cada HU sigue el formato narrativo Como/Quiero/Para, criterios de aceptación en Gherkin, "
    "reglas de negocio aplicables, Definition of Done y trazabilidad al issue de Jira."
)
add_paragraph(doc,
    "Las HU cumplen el criterio INVEST y aplican priorización MoSCoW. Para detalles operativos "
    "(estimaciones, asignaciones, sprints) consultar el Plan de Trabajo y el backlog de Jira TEC26.",
    italic=True
)

# Define HU data
hus_by_epic = {
    "Épica 1 — Análisis Técnico y Planeación (TEC26-1)": [
        ("TEC26-9",  "Como Arquitecto, quiero revisar la arquitectura actual de Preselección",
         "Identificar componentes, dependencias y patrones que afectan la migración a Azure PaaS",
         [("Revisión completa", "Dado que TecMilenio entregó la documentación técnica actualizada, "
                              "Cuando el Arquitecto revisa la arquitectura, Entonces se documenta el "
                              "inventario completo de componentes, frameworks y servicios"),
          ("Identificación de gaps", "Dado que la documentación tiene gaps, Cuando el Arquitecto los detecta, "
                                    "Entonces se levanta pregunta formal al PO en máximo 24 horas")],
         ["La revisión NO modifica la lógica de negocio del aplicativo",
          "Todo hallazgo se documenta en el RAID Log"],
         ["Documento de arquitectura actual entregado", "Inventario de componentes y versiones validado",
          "Preguntas al cliente cerradas", "Revisado por Scrum Master"],
         "Must Have", "3 puntos", "S1"),

        ("TEC26-10", "Como DevOps, quiero identificar las dependencias técnicas del aplicativo",
         "Definir los servicios Azure necesarios y validar viabilidad del Lift-and-Shift",
         [("Mapeo de dependencias", "Dado acceso al repositorio y documentación, Cuando el DevOps analiza el aplicativo, "
                                   "Entonces se entrega matriz de dependencias (framework, versión, propósito, equivalente Azure)"),
          ("Identificación de bloqueadores", "Dado una dependencia no compatible con Azure PaaS, Cuando se detecta, "
                                            "Entonces se registra en RAID con propuesta de mitigación")],
         ["Matriz de dependencias con equivalentes Azure"],
         ["Matriz entregada", "Equivalentes Azure mapeados", "Bloqueadores comunicados al PO"],
         "Must Have", "3 puntos", "S1"),

        ("TEC26-11", "Como Arquitecto, quiero validar requerimientos de ejecución en Azure",
         "Confirmar viabilidad del Lift-and-Shift sin reescritura",
         [("Validación de viabilidad", "Dado el inventario y dependencias, Cuando se evalúa contra App Service / "
                                       "Static Web Apps / Azure SQL, Entonces se entrega informe de viabilidad con riesgos")],
         [],
         ["Informe de viabilidad entregado", "Riesgos técnicos documentados en RAID", "Arquitectura objetivo aprobada"],
         "Must Have", "2 puntos", "S1"),

        ("TEC26-12", "Como Scrum Master, quiero documentar riesgos técnicos de la migración",
         "Gestionarlos proactivamente durante toda la ejecución del proyecto",
         [("Registro en RAID", "Dado los hallazgos del análisis, Cuando se identifica un riesgo, "
                              "Entonces se registra con probabilidad, impacto, severidad, mitigación y responsable")],
         [],
         ["RAID Log actualizado y compartido con PO", "Plan de mitigación para ALTO/CRÍTICO", "Revisión semanal calendarizada"],
         "Must Have", "2 puntos", "S1"),
    ],
    "Épica 2 — Preparación de Entornos Azure (TEC26-2)": [
        ("TEC26-13", "Como DevOps, quiero crear los Resource Groups en Azure",
         "Organizar y aislar los recursos del proyecto por ambiente",
         [("Provisión", "Dado acceso a la suscripción Azure, Cuando se ejecuta el provisionamiento, "
                       "Entonces existen RGs con nomenclatura rg-preseleccion-<ambiente>")],
         [],
         ["RGs creados con tags estándar", "Permisos RBAC configurados", "Validado por Arquitecto"],
         "Must Have", "2 puntos", "S1-S2"),

        ("TEC26-14", "Como DevOps, quiero provisionar App Services para el backend",
         "Habilitar el despliegue del aplicativo migrado",
         [("Provisión", "Dado los Resource Groups, Cuando se provisiona App Service Plan + App Service, "
                       "Entonces queda disponible un endpoint con runtime .NET Framework 4.8 y SSL")],
         [],
         ["App Service Plan + App Service No-Prod creados", "Slots configurados", "Auto-scaling configurado", "Health probe activo"],
         "Must Have", "3 puntos", "S2"),

        ("TEC26-15", "Como DevOps, quiero configurar Key Vault para gestión de secretos",
         "Centralizar la gestión de secretos sin hardcoding en código",
         [("Configuración", "Dado los Resource Groups y App Service, Cuando se crea Key Vault con Managed Identity, "
                           "Entonces los secretos se recuperan dinámicamente")],
         [],
         ["Key Vault creado con políticas RBAC", "Managed Identity habilitada", "Connection strings migradas", "Auditoría activa"],
         "Must Have", "2 puntos", "S2"),

        ("TEC26-16", "Como DevOps, quiero configurar Storage Account si el aplicativo lo requiere",
         "Soportar el manejo de archivos y datos no relacionales",
         [("Provisión condicional", "Dado que el análisis identificó uso de storage, Cuando se provisiona Storage, "
                                    "Entonces el aplicativo puede leer/escribir vía SAS o Managed Identity")],
         ["Sólo se provisiona si el análisis técnico lo identifica como necesario"],
         ["Storage Account creado", "Permisos configurados", "Conectividad validada"],
         "Should Have", "2 puntos", "S2"),

        ("TEC26-17", "Como DevOps, quiero configurar Azure Managed Redis si el aplicativo lo requiere",
         "Mantener la funcionalidad de cacheo en el entorno PaaS",
         [("Provisión condicional", "Dado uso de Redis identificado, Cuando se provisiona Azure Cache for Redis, "
                                    "Entonces el aplicativo conecta correctamente")],
         ["Sólo si el análisis técnico lo identifica como necesario"],
         ["Redis provisionado", "Cadena en Key Vault", "Validación cache hit/miss"],
         "Should Have", "2 puntos", "S2"),

        ("TEC26-18", "Como DevOps, quiero definir los ambientes No-Productivos y Productivo",
         "Soportar el ciclo completo de despliegue y promoción de cambios",
         [("Separación lógica", "Dado los recursos provisionados, Cuando se separan en ambientes, "
                                "Entonces cada uno tiene App Service, BD y secretos aislados")],
         ["Provisión productiva = responsabilidad TecMilenio"],
         ["Ambientes No-Prod y Prod definidos con tags", "Naming convention documentado"],
         "Must Have", "2 puntos", "S2"),
    ],
    "Épica 3 — Migración Lift-and-Shift del Aplicativo (TEC26-3)": [
        ("TEC26-19", "Como Dev, quiero migrar el backend .NET Framework + MVC a Azure App Service",
         "Habilitar la operación PaaS sin modificar la lógica de negocio",
         [("Despliegue exitoso", "Dado App Service provisionado, Cuando se publica el aplicativo, "
                                "Entonces inicia correctamente y responde a healthcheck"),
          ("Sin cambios de lógica", "Dado el código migrado, Cuando se compara con el original, "
                                   "Entonces la lógica de negocio se mantiene sin cambios")],
         ["SIN modificación de lógica de negocio",
          "Upgrade dentro de misma familia: .NET 4.7.x → 4.8.1 PERMITIDO",
          "Reescritura o cambio de framework: PROHIBIDO"],
         ["Backend desplegado en App Service", "Healthcheck pasa", "Logs en Application Insights", "Code review por Arquitecto"],
         "Must Have", "5 puntos", "S2"),

        ("TEC26-20", "Como Dev, quiero migrar el frontend AngularJS a Static Web Apps",
         "Servir el frontend desde un servicio PaaS optimizado para SPAs",
         [("Despliegue", "Dado frontend AngularJS, Cuando se despliega en Static Web Apps, "
                        "Entonces carga correctamente y se conecta con el backend")],
         ["Update permitido: AngularJS 1.5 → 1.8.x (misma familia)"],
         ["Frontend desplegado", "Llamadas API funcionan vía CORS", "Tiempos de carga ≤ versión anterior"],
         "Should Have", "3 puntos", "S2"),

        ("TEC26-21", "Como Dev, quiero configurar variables de entorno y secretos en App Service",
         "Que opere correctamente sin hardcoding de credenciales en código",
         [("Configuración", "Dado backend desplegado y Key Vault, Cuando se cargan App Settings desde Key Vault, "
                           "Entonces el aplicativo arranca sin errores")],
         [],
         ["App Settings configuradas", "Secretos desde Key Vault", "Validado que NO existen secretos en repo"],
         "Must Have", "2 puntos", "S2"),

        ("TEC26-22", "Como Dev, quiero configurar las cadenas de conexión a base de datos",
         "Garantizar la conectividad con la base de datos migrada",
         [("Conexión", "Dado Azure SQL provisionado, Cuando el aplicativo intenta conectar, "
                      "Entonces es exitoso con cadenas en Key Vault")],
         [],
         ["Connection strings en Key Vault", "App Service consume vía Managed Identity", "Validado con query simple"],
         "Must Have", "2 puntos", "S2"),

        ("TEC26-23", "Como Dev, quiero validar el manejo de archivos y rutas en cloud",
         "Garantizar que el aplicativo funcione idénticamente que en IaaS",
         [("Validación I/O", "Dado el aplicativo en App Service, Cuando se ejecutan operaciones I/O, "
                            "Entonces el comportamiento es equivalente al ambiente IaaS")],
         [],
         ["Rutas validadas (System.IO, Server.MapPath, ~/)", "Archivos persistentes en Storage si aplica", "Smoke test I/O pasa"],
         "Must Have", "3 puntos", "S2"),
    ],
    "Épica 4 — Migración de Bases de Datos MS SQL Server (TEC26-4)": [
        ("TEC26-24", "Como DevOps, quiero migrar esquemas y datos MS SQL hacia Azure PaaS",
         "Soportar la operación del aplicativo migrado",
         [("Migración completa", "Dado acceso a BD origen y destino, Cuando se ejecuta migración (DMS/BACPAC), "
                                 "Entonces todos los esquemas y datos están en Azure SQL"),
          ("Conteos coinciden", "Dado origen y destino, Cuando se comparan conteos, "
                                "Entonces coinciden 100%")],
         ["Backup pre-migración obligatorio", "Plan de rollback documentado y probado"],
         ["Backup tomado", "Migración ejecutada", "Conteos validados", "Rollback probado"],
         "Must Have", "5 puntos", "S2 — RIESGO ALTO (R-01)"),

        ("TEC26-25", "Como DevOps, quiero configurar la conectividad del aplicativo con la BD migrada",
         "Restaurar la operación end-to-end del aplicativo en cloud",
         [("Conectividad", "Dado Azure SQL con datos y App Service desplegado, Cuando se configura conexión, "
                          "Entonces el aplicativo realiza consultas exitosamente")],
         [],
         ["Firewall configurado", "Connection strings probadas en cada ambiente", "Auditoría activa"],
         "Must Have", "2 puntos", "S2"),

        ("TEC26-26", "Como QA, quiero validar completitud de la información migrada",
         "Asegurar que no hay pérdida de datos respecto al origen",
         [("Validación", "Dado BD origen y migrada, Cuando se ejecutan conteos/checksums/muestreo, "
                        "Entonces todos los registros y tablas coinciden 100%")],
         [],
         ["Conteos validados", "Checksums validados", "Muestreo >30 registros por tabla", "Reporte entregado"],
         "Must Have", "3 puntos", "S2"),

        ("TEC26-27", "Como QA, quiero ejecutar smoke tests con datos migrados",
         "Validar que la operación end-to-end funciona con datos reales",
         [("Smoke tests", "Dado el aplicativo conectado a Azure SQL, Cuando se ejecuta el set, "
                         "Entonces todos los escenarios pasan sin errores críticos")],
         [],
         ["Set definido y ejecutado", "Resultados documentados", "Issues abiertos para fallas"],
         "Must Have", "2 puntos", "S2"),
    ],
    "Épica 5 — Pipelines CI/CD (TEC26-5)": [
        ("TEC26-28", "Como DevOps, quiero configurar pipelines de despliegue automatizado",
         "Automatizar y estandarizar los despliegues a los ambientes",
         [("Pipeline funcional", "Dado el repo con código migrado, Cuando se hace push, "
                                "Entonces el pipeline ejecuta build/tests/deploy")],
         [],
         ["Pipeline build+deploy a No-Prod funcional", "Pipeline a Prod con aprobación manual", "Documentación entregada"],
         "Must Have", "5 puntos", "S2-S3"),

        ("TEC26-29", "Como DevOps, quiero integrar pipelines con repositorios GitHub existentes",
         "No romper el flujo de trabajo del equipo cliente",
         [("Integración", "Dado acceso a repos GitHub TecMilenio, Cuando se configura integración, "
                         "Entonces el pipeline reacciona a eventos del repo")],
         [],
         ["Webhooks/triggers configurados", "Secretos GitHub configurados", "PR validation automática"],
         "Must Have", "2 puntos", "S2-S3"),

        ("TEC26-30", "Como DevOps, quiero definir los flujos de deployment hacia ambientes",
         "Controlar la promoción de cambios entre ambientes",
         [("Estrategia documentada", "Dado los pipelines, Cuando se establece la estrategia, "
                                     "Entonces queda documentada y validada con el equipo")],
         [],
         ["Diagrama de flujo entregado", "Gates de promoción definidos", "Runbook entregado"],
         "Must Have", "2 puntos", "S3"),
    ],
    "Épica 6 — Atención de 13 Hallazgos de Seguridad (TEC26-6)": [
        ("TEC26-31", "Como SecOps, quiero atender los 13 hallazgos de seguridad de TecMilenio",
         "Cumplir con los estándares de seguridad acordados antes del Go-Live",
         [("Mitigación", "Dado los 13 hallazgos confirmados, Cuando se ejecutan mitigaciones, "
                        "Entonces cada uno queda atendido y documentado en su subtarea (TEC26-32 a TEC26-44)")],
         ["13 hallazgos en alcance — adicionales requieren change order",
          "Modernizr y Html5Shiv son librerías deprecadas — pueden requerir aceptación de riesgo"],
         ["13 subtareas cerradas", "Evidencia documental por mitigación", "Revisión Arquitecto + SecOps"],
         "Must Have", "8 puntos", "S3"),

        ("TEC26-45", "Como SecOps, quiero ejecutar el reescaneo de los 13 hallazgos atendidos",
         "Validar formalmente que las mitigaciones aplicadas son efectivas",
         [("Reescaneo", "Dado las 13 subtareas cerradas, Cuando se ejecuta el reescaneo, "
                       "Entonces ningún hallazgo crítico permanece abierto")],
         [],
         ["Reescaneo ejecutado", "Reporte comparativo (antes/después)", "Validación entregada a TecMilenio"],
         "Must Have", "2 puntos", "S3"),

        ("TEC26-46", "Como QA, quiero validar el aplicativo posterior a los cambios de seguridad",
         "Asegurar que no se introdujeron regresiones funcionales",
         [("Regresión", "Dado las 13 mitigaciones, Cuando se ejecuta validación funcional, "
                       "Entonces los flujos críticos siguen funcionando")],
         [],
         ["Set de regresión ejecutado", "Sin defectos críticos abiertos", "Reporte entregado"],
         "Must Have", "2 puntos", "S3"),
    ],
    "Épica 7 — Pruebas y UAT (TEC26-7)": [
        ("TEC26-47", "Como QA, quiero ejecutar smoke tests del aplicativo migrado",
         "Validar operación, conectividad e integraciones existentes",
         [("Smoke tests", "Dado el aplicativo en No-Prod con datos migrados, Cuando se ejecutan smoke tests, "
                         "Entonces todos los flujos críticos pasan sin hallazgos bloqueantes")],
         [],
         ["Smoke tests ejecutados", "Resultados con evidencias", "Defectos críticos resueltos"],
         "Must Have", "3 puntos", "S3"),

        ("TEC26-48", "Como QA Mobiik, quiero acompañar las pruebas UAT del cliente",
         "Resolver dudas técnicas, replicar incidencias y agilizar la aceptación",
         [("Acompañamiento", "Dado UAT programado (3 días), Cuando el cliente reporta una incidencia, "
                            "Entonces se replica, clasifica y documenta en máximo 1 día hábil")],
         [],
         ["Acompañamiento diario durante 3 días", "Bitácora de incidencias registradas", "Comunicación con PO"],
         "Must Have", "5 puntos", "S3"),

        ("TEC26-49", "Como Dev, quiero corregir las incidencias detectadas durante UAT",
         "Cumplir con los criterios de aceptación del cliente",
         [("Correcciones", "Dado incidencias clasificadas por severidad, Cuando se corrigen las críticas/altas, "
                          "Entonces el aplicativo está listo para aceptación final")],
         ["Sólo se corrigen incidencias derivadas de la migración",
          "Bugs preexistentes → change order"],
         ["Incidencias críticas/altas resueltas", "Validación cruzada con QA", "Documentación actualizada"],
         "Must Have", "5 puntos", "S3"),

        ("TEC26-50", "Como Scrum Master, quiero obtener firma de aceptación del cliente",
         "Formalizar el cierre de la fase de pruebas previo al Go-Live",
         [("Sign-off", "Dado UAT ejecutado y defectos resueltos, Cuando se presenta el reporte final, "
                      "Entonces se firma el documento de aceptación")],
         [],
         ["Documento firmado por PO TecMilenio", "Archivo entregable adjunto", "Comunicación al equipo"],
         "Must Have", "1 punto", "S3"),
    ],
    "Épica 8 — Go-Live y Cierre del Proyecto (TEC26-8)": [
        ("TEC26-51", "Como DevOps, quiero ejecutar el despliegue productivo",
         "Liberar la solución a los usuarios finales de TecMilenio",
         [("Despliegue Prod", "Dado infraestructura productiva (TecMilenio) y aceptación firmada, "
                              "Cuando se ejecuta el pipeline a Prod, Entonces el aplicativo está operativo en producción")],
         ["Infraestructura productiva = responsabilidad TecMilenio"],
         ["Despliegue ejecutado", "Smoke test post-deployment OK", "Monitoreo activo"],
         "Must Have", "3 puntos", "Cierre"),

        ("TEC26-52", "Como Scrum Master, quiero acompañar al cliente post-liberación",
         "Atender incidencias inmediatas y asegurar estabilización productiva",
         [("Acompañamiento inicial", "Dado el aplicativo en producción, Cuando se reportan incidentes, "
                                     "Entonces se atienden con SLA acordado")],
         ["Garantía de 30 días naturales post-aceptación final",
          "Cobertura sólo dentro de horario laboral L-V 9-18"],
         ["Bitácora de soporte registrada", "Hand-off a soporte interno TecMilenio"],
         "Must Have", "2 puntos", "Cierre"),

        ("TEC26-53", "Como Scrum Master, quiero entregar la Carta de Cierre del proyecto",
         "Formalizar el cierre del proyecto y liberar el Hito de Facturación 2 (70% — $153,153 MXN)",
         [("Entrega final", "Dado el aplicativo en producción y soporte completado, "
                            "Cuando se entrega Carta de Cierre + Memoria Técnica, "
                            "Entonces TecMilenio firma la conclusión del proyecto")],
         [],
         ["Carta de cierre firmada por TecMilenio", "Memoria técnica entregada", "Facturación Hito 2 liberada"],
         "Must Have", "1 punto", "Cierre"),
    ],
}

# Render HUs
for epic_title, hus in hus_by_epic.items():
    add_h2(doc, epic_title)
    for hu in hus:
        jira_id, summary, purpose, acs, rules, dod, prio, est, sprint = hu

        # HU header card (single-cell table with dark bg)
        hdr_table = doc.add_table(rows=1, cols=1)
        hdr_table.autofit = False
        hdr_cell = hdr_table.rows[0].cells[0]
        hdr_cell.width = Inches(6.5)
        set_cell_bg(hdr_cell, "0A0A0A")
        hdr_cell.text = ""
        p_id = hdr_cell.paragraphs[0]
        r_id = p_id.add_run(jira_id)
        r_id.font.name = FONT; r_id.font.size = Pt(10); r_id.font.bold = True; r_id.font.color.rgb = LIME
        p_id.add_run("   ·   ").font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r_sp = p_id.add_run(f"Sprint {sprint}  ·  Prioridad: {prio}  ·  Estimación: {est}")
        r_sp.font.name = FONT; r_sp.font.size = Pt(9); r_sp.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        p_sum = hdr_cell.add_paragraph()
        r_sum = p_sum.add_run(summary)
        r_sum.font.name = FONT; r_sum.font.size = Pt(13); r_sum.font.bold = True; r_sum.font.color.rgb = WHITE

        # Narrativa
        p_narr = doc.add_paragraph()
        p_narr.paragraph_format.space_before = Pt(6)
        r_narr = p_narr.add_run(f"Para que ")
        r_narr.font.name = FONT; r_narr.font.size = Pt(11)
        r_narr2 = p_narr.add_run(purpose + ".")
        r_narr2.font.name = FONT; r_narr2.font.size = Pt(11); r_narr2.font.italic = True

        # Criterios de aceptación
        add_h3(doc, "Criterios de Aceptación (Gherkin)")
        for ac_name, ac_text in acs:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            r1 = p.add_run(f"Escenario — {ac_name}: ")
            r1.font.name = FONT; r1.font.size = Pt(10.5); r1.font.bold = True; r1.font.color.rgb = DARK_GREY
            r2 = p.add_run(ac_text)
            r2.font.name = FONT; r2.font.size = Pt(10.5)

        # Reglas de Negocio
        if rules:
            add_h3(doc, "Reglas de Negocio")
            for rule in rules:
                add_bullet(doc, rule, bold_prefix="• ")

        # DoD
        add_h3(doc, "Definition of Done")
        for d in dod:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(f"☐  {d}")
            r.font.name = FONT; r.font.size = Pt(10.5)

        # Spacer
        doc.add_paragraph()

doc.add_page_break()


# ═══ 7. MATRIZ DE TRAZABILIDAD ═════════════════════════════════════
add_heading_bar(doc, "7. Matriz de Trazabilidad Comercial ↔ Funcional", size=18)
add_paragraph(doc,
    "Esta matriz garantiza que cada requerimiento de la propuesta comercial está cubierto por una o más HUs en Jira."
)

matriz = [
    ["Análisis técnico inicial",                  "Sec. Alcance > Análisis técnico",   "TEC26-1",  "S1",      "Hito 1"],
    ["Preparación de entornos Azure",              "Sec. Alcance > Preparación",         "TEC26-2",  "S1-S2",   "Hito 2"],
    ["Migración Lift-and-Shift",                   "Sec. Alcance > Migración",           "TEC26-3",  "S2",      "Hito 2"],
    ["Migración BD MS SQL Server",                  "Sec. Alcance > BD",                  "TEC26-4",  "S2",      "Hito 2"],
    ["Pipelines CI/CD",                            "Sec. Alcance > DevOps básico",       "TEC26-5",  "S2-S3",   "Hito 2"],
    ["Atención de 13 hallazgos de seguridad",      "Sec. Alcance > Hallazgos",          "TEC26-6",  "S3",      "Hito 2"],
    ["Pruebas (smoke + UAT)",                      "Sec. Alcance > Pruebas y validación","TEC26-7", "S3",      "Hito 2"],
    ["Puesta en operación (Go-Live)",              "Sec. Alcance > Puesta en operación", "TEC26-8",  "Cierre",  "Hito 2"],
]
add_branded_table(doc,
    ["Requerimiento Comercial", "Sección Propuesta", "Épica Jira", "Sprint", "Hito Facturación"],
    matriz,
    col_widths=[Inches(2.0), Inches(1.5), Inches(0.9), Inches(0.9), Inches(1.2)]
)


# ═══ 8. SUPUESTOS Y RESTRICCIONES ══════════════════════════════════
add_heading_bar(doc, "8. Supuestos y Restricciones", size=18)
add_paragraph(doc, "Los siguientes supuestos rigen el alcance funcional del proyecto. Si alguno no se cumple, "
                  "se requiere change order para ajustar costo y/o cronograma.")
sup_data = [
    ["A-01", "Cliente validó funcionalmente el aplicativo y no tiene bugs conocidos"],
    ["A-02", "Documentación técnica del aplicativo disponible y actualizada"],
    ["A-03", "Acceso a repositorios GitHub y pipelines existentes provisto por TecMilenio"],
    ["A-04", "Se conserva la familia tecnológica actual (sin upgrade a familias superiores)"],
    ["A-05", "Provisión y configuración de infraestructura productiva la realiza TecMilenio"],
    ["A-06", "Validación funcional completa post-liberación es responsabilidad del cliente"],
    ["A-07", "Sin pruebas de pentesting ni pruebas de carga (fuera de alcance)"],
    ["A-08", "Costos de infraestructura Azure los cubre TecMilenio"],
    ["A-09", "Garantía técnica de 30 días naturales post-aceptación"],
    ["A-10", "Servicios prestados de manera remota (sin viáticos contemplados)"],
]
add_branded_table(doc, ["ID", "Supuesto"], sup_data, col_widths=[Inches(0.6), Inches(5.9)])


# ═══ 9. FUERA DE ALCANCE ═══════════════════════════════════════════
add_heading_bar(doc, "9. Fuera de Alcance", size=18)
add_paragraph(doc, "Los siguientes elementos NO forman parte del alcance del presente proyecto y requieren contratación adicional:")
out_of_scope = [
    "Pruebas de rendimiento o pruebas de carga sobre el aplicativo migrado",
    "Pentesting formal de seguridad",
    "Capacitación de usuarios finales sobre el aplicativo",
    "Soporte y capacitación para uso del aplicativo a nivel usuario final",
    "Costo de viáticos y alojamiento",
    "Costos de infraestructura Azure (desarrollo, pruebas y producción)",
    "Licenciamiento, hardware o software adicional",
    "Carga de información adicional a la migración de base de datos",
    "Integración con sistemas adicionales no mencionados en la propuesta",
    "Creación o edición de contenido",
    "Gestión funcional de usuarios, permisos y roles fuera de los requeridos",
    "Mantenimiento Post-Migración",
    "Optimización Post-Migración",
    "Desarrollo de Nuevas Funcionalidades",
    "Modificación o soporte a integraciones existentes fuera de los ajustes mínimos",
    "Migración de Datos no Relacionados",
    "Gestión de Cambio Organizacional",
    "Upgrade de tecnologías a familias superiores a la base actual",
    "Atención de hallazgos adicionales a los 13 identificados durante la estimación inicial",
]
for item in out_of_scope:
    add_bullet(doc, item, bold_prefix="✗ ")


# ═══ 10. ANEXOS ════════════════════════════════════════════════════
add_heading_bar(doc, "10. Anexos", size=18)
add_h3(doc, "Anexo A — Listado de 13 Hallazgos de Seguridad en alcance")
add_branded_table(doc,
    ["ID", "Hallazgo", "Subtarea Jira", "Severidad"],
    [
        ["1",  "Out-of-date Version (jQuery)",       "TEC26-32", "Media"],
        ["2",  "Out-of-date Version (Bootstrap)",    "TEC26-33", "Media"],
        ["3",  "Internal Server Error",               "TEC26-34", "Alta"],
        ["4",  "Cookie Not Marked as HttpOnly",       "TEC26-35", "Alta"],
        ["5",  "Version Disclosure (ASP.NET)",        "TEC26-36", "Baja"],
        ["6",  "Version Disclosure (IIS)",            "TEC26-37", "Baja"],
        ["7",  "Version Disclosure (ASP.NET MVC)",    "TEC26-38", "Baja"],
        ["8",  "Cross-site Request Forgery (CSRF)",   "TEC26-39", "Crítica"],
        ["9",  "Version Disclosure (jQuery)",         "TEC26-40", "Baja"],
        ["10", "Version Disclosure (Bootstrap.js)",   "TEC26-41", "Baja"],
        ["11", "Version Disclosure (Modernizr)",      "TEC26-42", "Baja"],
        ["12", "Version Disclosure (Html5Shiv)",      "TEC26-43", "Baja"],
        ["13", "Missing X-Content-Type-Options",      "TEC26-44", "Media"],
    ],
    col_widths=[Inches(0.5), Inches(3.0), Inches(1.3), Inches(1.0)]
)

add_h3(doc, "Anexo B — Documentos complementarios")
add_bullet(doc, "Propuesta comercial: Migración Técnica Preselección — Mobiik × TecMilenio (20-may-2026)", bold_prefix="📄 ")
add_bullet(doc, "RAID Log: RAID Log - Migración Preselección.xlsx", bold_prefix="📄 ")
add_bullet(doc, "Plan de Trabajo: Plan de Trabajo - Migración Preselección.xlsx", bold_prefix="📄 ")
add_bullet(doc, "PPT Kickoff: KO TECMILENIO 2.6 - Mobiik Standard.pptx", bold_prefix="📄 ")
add_bullet(doc, "Backlog Jira: TEC26 en mobiik-jira.atlassian.net", bold_prefix="🔗 ")


# ═══ 11. SIGN-OFF ══════════════════════════════════════════════════
doc.add_page_break()
add_heading_bar(doc, "11. Aceptación y Sign-off", size=18)
add_paragraph(doc,
    "Mediante la firma del presente documento, ambas partes manifiestan su conformidad con el alcance funcional "
    "definido en las Historias de Usuario, los supuestos establecidos y las restricciones aplicables. La aceptación "
    "formaliza el inicio del desarrollo y sirve como referencia para las pruebas UAT y la liberación productiva."
)

doc.add_paragraph()
doc.add_paragraph()

# Two-column signature block
sig_table = doc.add_table(rows=1, cols=2)
sig_table.autofit = False
for col in sig_table.columns:
    col.width = Inches(3.0)

# TecMilenio
left = sig_table.rows[0].cells[0]
left.text = ""
p1 = left.paragraphs[0]
r1 = p1.add_run("POR TECMILENIO")
r1.font.name = FONT; r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = LIME
left.add_paragraph()
for label in ["Nombre:", "Cargo:", "Firma:", "Fecha:"]:
    p = left.add_paragraph()
    r = p.add_run(f"{label}  ____________________________")
    r.font.name = FONT; r.font.size = Pt(10)

# Mobiik
right = sig_table.rows[0].cells[1]
right.text = ""
p2 = right.paragraphs[0]
r2 = p2.add_run("POR MOBIIK")
r2.font.name = FONT; r2.font.size = Pt(11); r2.font.bold = True; r2.font.color.rgb = LIME
right.add_paragraph()
for label in ["Nombre:", "Cargo:", "Firma:", "Fecha:"]:
    p = right.add_paragraph()
    r = p.add_run(f"{label}  ____________________________")
    r.font.name = FONT; r.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

# Final footer
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_final = p_final.add_run("© 2026 Mobiik   |   coding for the future   |   Documento Confidencial")
r_final.font.name = FONT; r_final.font.size = Pt(9); r_final.font.italic = True; r_final.font.color.rgb = GREY


# Save
out = "./entregables/Historias de Usuario - Migración Preselección.docx"
doc.save(out)
print(f"✓ Saved: {out}")
