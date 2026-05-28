"""
Minuta de Reunión — Mobiik
Acta de seguimiento de proyecto / sesión de trabajo.
Estándar visual Mobiik: encabezado branded, acento verde lima #AADC1E,
tablas con header oscuro. Parafrasea, no transcribe.

Basado en la plantilla corporativa Mobiik_Template_Minuta.docx
Dependencia:  pip3 install --user python-docx
Uso:          python3 build_minuta_docx.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── DATOS DE LA MINUTA (editar) ─────────────────────────────────────
PROYECTO     = "[Nombre del proyecto]"
CLIENTE      = "[Cliente]"
FECHA        = "[DD/MM/AAAA]"
HORA         = "[HH:MM – HH:MM]"
LUGAR_MEDIO  = "[Teams / Presencial]"
ELABORO      = "[Nombre]"
OBJETIVO     = "[Describir en una o dos líneas el propósito de la sesión.]"

# (Nombre, Rol / Área, Organización)
ASISTENTES = [
    ("[Nombre]", "[Rol / Área]", "[Mobiik / Cliente]"),
]

# Temas tratados — resúmenes parafraseados, NO transcripción literal
TEMAS = [
    "[Tema 1 — resumen del punto discutido y conclusión]",
    "[Tema 2 — ...]",
    "[Tema 3 — ...]",
]

ACUERDOS = [
    "[Acuerdo 1]",
    "[Acuerdo 2]",
]

# (Acción / Compromiso, Responsable, Fecha límite, Estatus)
COMPROMISOS = [
    ("[Acción / compromiso]", "[Responsable]", "[DD/MM/AAAA]", "Abierto"),
]

# Pendientes para la próxima sesión (usar [POR CONFIRMAR] si hay ambigüedad)
PENDIENTES = [
    "[Pendiente 1]",
]

# Riesgos o bloqueos detectados en la sesión ([] si ninguno)
RIESGOS = [
    "[Riesgo / bloqueo detectado en la sesión, o dejar la lista vacía]",
]

PROXIMA_FECHA = "[DD/MM/AAAA, HH:MM]"
PROXIMA_TEMAS = "[____]"

OUT = "./entregables/Minuta - <proyecto> - <fecha>.docx"
# ─────────────────────────────────────────────────────────────────────

# ─── Brand ───────────────────────────────────────────────────────────
LIME      = RGBColor(0xAA, 0xDC, 0x1E)
BLACK     = RGBColor(0x0A, 0x0A, 0x0A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY      = RGBColor(0x66, 0x66, 0x66)
ORANGE    = RGBColor(0xFF, 0xA0, 0x40)
FONT      = "Arial"

doc = Document()

sec = doc.sections[0]
sec.page_height = Inches(11)
sec.page_width  = Inches(8.5)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(11)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_heading_bar(text, *, size=16, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'AADC1E')
    pBdr.append(left)
    p_pr.append(pBdr)
    return p

def add_paragraph(text, *, italic=False, color=None, size=11, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(11)
    return p

def add_branded_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr[i]
        set_cell_bg(c, "0A0A0A"); set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = FONT; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            c = cells[ci]
            set_cell_bg(c, "F5F5F5" if ri % 2 == 0 else "FFFFFF"); set_cell_border(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = FONT; r.font.size = Pt(10)
            # Resalta marcas de ambigüedad
            if "[POR CONFIRMAR]" in str(val):
                r.font.color.rgb = ORANGE; r.font.bold = True
    return table

def add_kv_grid(rows):
    """Datos generales en grid 2 columnas (etiqueta oscura + valor)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for ri, (k, v) in enumerate(rows):
        kc, vc = table.rows[ri].cells
        kc.width = Inches(2.0); vc.width = Inches(4.5)
        set_cell_bg(kc, "0A0A0A"); set_cell_border(kc)
        set_cell_bg(vc, "F5F5F5" if ri % 2 == 0 else "FFFFFF"); set_cell_border(vc)
        kc.vertical_alignment = vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        kc.text = ""; vc.text = ""
        rk = kc.paragraphs[0].add_run(k)
        rk.font.name = FONT; rk.font.size = Pt(10); rk.font.bold = True; rk.font.color.rgb = WHITE
        rv = vc.paragraphs[0].add_run(v)
        rv.font.name = FONT; rv.font.size = Pt(10.5)
    return table


# ═══════════════════════════════════════════════════════════════════
# ENCABEZADO
# ═══════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
rt = title.add_run("Minuta de Reunión")
rt.font.name = FONT; rt.font.size = Pt(24); rt.font.bold = True; rt.font.color.rgb = BLACK
p_pr = title._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
left = OxmlElement('w:left')
left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24')
left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'AADC1E')
pBdr.append(left); p_pr.append(pBdr)

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(16)
rs = sub.add_run("Acta de seguimiento de proyecto / sesión de trabajo")
rs.font.name = FONT; rs.font.size = Pt(11); rs.font.italic = True; rs.font.color.rgb = GREY

# Datos generales
add_kv_grid([
    ("Proyecto", PROYECTO),
    ("Cliente", CLIENTE),
    ("Fecha", FECHA),
    ("Hora", HORA),
    ("Lugar / Medio", LUGAR_MEDIO),
    ("Elaboró", ELABORO),
])

# ═══════════════════════════════════════════════════════════════════
# OBJETIVO
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Objetivo de la reunión")
add_paragraph(OBJETIVO, after=8)

# ═══════════════════════════════════════════════════════════════════
# ASISTENTES
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Asistentes")
add_branded_table(
    ["Nombre", "Rol / Área", "Organización"],
    ASISTENTES,
    col_widths=[Inches(2.4), Inches(2.4), Inches(1.7)],
)

# ═══════════════════════════════════════════════════════════════════
# TEMAS TRATADOS
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Temas tratados")
for t in TEMAS:
    add_bullet(t)

# ═══════════════════════════════════════════════════════════════════
# ACUERDOS
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Acuerdos y decisiones")
for a in ACUERDOS:
    add_bullet(a)

# ═══════════════════════════════════════════════════════════════════
# COMPROMISOS Y RESPONSABLES
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Compromisos y responsables")
add_branded_table(
    ["Acción / Compromiso", "Responsable", "Fecha límite", "Estatus"],
    COMPROMISOS,
    col_widths=[Inches(3.0), Inches(1.5), Inches(1.2), Inches(0.9)],
)

# ═══════════════════════════════════════════════════════════════════
# PENDIENTES PARA LA PRÓXIMA SESIÓN
# ═══════════════════════════════════════════════════════════════════
if PENDIENTES:
    add_heading_bar("Pendientes para la próxima sesión")
    for p in PENDIENTES:
        add_bullet(p)

# ═══════════════════════════════════════════════════════════════════
# RIESGOS / BLOQUEOS
# ═══════════════════════════════════════════════════════════════════
if RIESGOS:
    add_heading_bar("Riesgos y bloqueos")
    for r in RIESGOS:
        add_bullet(r)

# ═══════════════════════════════════════════════════════════════════
# PRÓXIMA REUNIÓN
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Próxima reunión")
add_paragraph(
    f"Fecha y hora: {PROXIMA_FECHA}   |   Temas a tratar: {PROXIMA_TEMAS}", after=20)

# ═══════════════════════════════════════════════════════════════════
# FIRMAS
# ═══════════════════════════════════════════════════════════════════
sig = doc.add_table(rows=1, cols=2)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
for col in sig.columns:
    col.width = Inches(3.0)
for cell, label in zip(sig.rows[0].cells, ["Por Mobiik", "Por el Cliente"]):
    cell.text = ""
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(24)
    r0 = p0.add_run("_______________________________")
    r0.font.name = FONT; r0.font.size = Pt(11)
    p1 = cell.add_paragraph()
    r1 = p1.add_run(label)
    r1.font.name = FONT; r1.font.size = Pt(10); r1.font.bold = True; r1.font.color.rgb = GREY

# Pie de página
doc.add_paragraph()
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run("© 2026 Mobiik   |   coding for the future   |   Documento Confidencial")
rf.font.name = FONT; rf.font.size = Pt(9); rf.font.italic = True; rf.font.color.rgb = GREY


# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
import os
os.makedirs(os.path.dirname(OUT) or ".", exist_ok=True)
doc.save(OUT)
print(f"✓ Saved: {OUT}")
