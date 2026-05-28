"""
Carta de Cierre de Proyecto — Mobiik
Formaliza la conclusión y aceptación de entregables ante el cliente.
Estándar visual Mobiik: portada/encabezado branded, acento verde lima #AADC1E,
tablas con header oscuro, tono cordial y relationship-forward.

Basado en la plantilla corporativa Mobiik_Template_Carta_Cierre.docx
Dependencia:  pip3 install --user python-docx
Uso:          python3 build_carta_cierre_docx.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── DATOS DE LA CARTA (editar) ──────────────────────────────────────
CIUDAD        = "Ciudad de México"
FECHA         = "[DD] de [mes] de [AAAA]"
DESTINATARIO  = "[Nombre del destinatario]"
CARGO_DEST    = "[Cargo]"
EMPRESA_DEST  = "[Empresa / Área]"
SALUDO_NOMBRE = "[Nombre]"

PROYECTO      = "[Nombre del proyecto]"
CLIENTE       = "[Cliente]"
PERIODO       = "[Inicio] – [Fin]"
PATROCINADOR  = "[Nombre]"
CONTRATO_OC   = "[Referencia]"

# Tipo de cierre: "proyecto", "fase" o "entregable"
TIPO_CIERRE   = "proyecto"

OBJETIVOS = [
    "[Objetivo / resultado 1]",
    "[Objetivo / resultado 2]",
    "[Objetivo / resultado 3]",
]

# (Entregable, Fecha de entrega, Aceptado Sí/No)
ENTREGABLES = [
    ("[Entregable 1]", "[DD/MM/AAAA]", "Sí"),
    ("[Entregable 2]", "[DD/MM/AAAA]", "Sí"),
]

PENDIENTES = "[Indicar si existen puntos de garantía, soporte post-implementación, o anotar “Ninguno”.]"

# Firmantes
FIRMA_MOBIIK  = ("[Nombre]", "[Cargo] — Mobiik")
FIRMA_CLIENTE = ("[Nombre]", "[Cargo] — [Cliente]")

OUT = "./entregables/Carta de Cierre - <proyecto>.docx"
# ─────────────────────────────────────────────────────────────────────

# ─── Brand ───────────────────────────────────────────────────────────
LIME      = RGBColor(0xAA, 0xDC, 0x1E)
BLACK     = RGBColor(0x0A, 0x0A, 0x0A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY      = RGBColor(0x66, 0x66, 0x66)
DARK_GREY = RGBColor(0x25, 0x25, 0x25)
FONT      = "Arial"

doc = Document()

sec = doc.sections[0]
sec.page_height = Inches(11)
sec.page_width  = Inches(8.5)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(11)


# ─── Helpers (idénticos al estándar Mobiik) ──────────────────────────
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_heading_bar(text, *, size=18, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'AADC1E')
    pBdr.append(left)
    p_pr.append(pBdr)
    return p

def add_paragraph(text, *, italic=False, color=None, size=11, after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(11)
    return p

def add_kv_table(rows):
    """Tabla clave/valor: etiqueta oscura a la izquierda, valor a la derecha."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for ri, (k, v) in enumerate(rows):
        kc, vc = table.rows[ri].cells
        kc.width = Inches(2.3); vc.width = Inches(4.2)
        set_cell_bg(kc, "0A0A0A"); set_cell_border(kc)
        set_cell_bg(vc, "F5F5F5" if ri % 2 == 0 else "FFFFFF"); set_cell_border(vc)
        kc.vertical_alignment = vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        kc.text = ""; vc.text = ""
        rk = kc.paragraphs[0].add_run(k)
        rk.font.name = FONT; rk.font.size = Pt(10); rk.font.bold = True; rk.font.color.rgb = WHITE
        rv = vc.paragraphs[0].add_run(v)
        rv.font.name = FONT; rv.font.size = Pt(10.5)
    return table

def add_branded_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr[i]
        set_cell_bg(c, "0A0A0A"); set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = FONT; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            c = cells[ci]
            set_cell_bg(c, "F5F5F5" if ri % 2 == 0 else "FFFFFF"); set_cell_border(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = FONT; r.font.size = Pt(10)
    return table

def add_signature_block(firmante, etiqueta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("_______________________________")
    r.font.name = FONT; r.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(firmante)
    r2.font.name = FONT; r2.font.size = Pt(11); r2.font.bold = True
    p3 = doc.add_paragraph()
    r3 = p3.add_run(etiqueta)
    r3.font.name = FONT; r3.font.size = Pt(10); r3.font.color.rgb = GREY


# ═══════════════════════════════════════════════════════════════════
# ENCABEZADO
# ═══════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
rt = title.add_run("Carta de Cierre de Proyecto")
rt.font.name = FONT; rt.font.size = Pt(24); rt.font.bold = True; rt.font.color.rgb = BLACK
# Barra verde lima bajo el título
p_pr = title._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
left = OxmlElement('w:left')
left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24')
left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'AADC1E')
pBdr.append(left); p_pr.append(pBdr)

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(18)
rs = sub.add_run("Formalización de la conclusión y aceptación de entregables")
rs.font.name = FONT; rs.font.size = Pt(11); rs.font.italic = True; rs.font.color.rgb = GREY

# Lugar y fecha
add_paragraph(f"{CIUDAD}, a {FECHA}", align=WD_ALIGN_PARAGRAPH.RIGHT, after=14)

# Destinatario
add_paragraph(DESTINATARIO, after=0)
add_paragraph(CARGO_DEST, after=0)
add_paragraph(EMPRESA_DEST, after=14)

add_paragraph(f"Estimado(a) {SALUDO_NOMBRE}:", after=10)

# Cuerpo introductorio
add_paragraph(
    f"Por medio de la presente, Mobiik hace constar la conclusión formal del "
    f"{TIPO_CIERRE} “{PROYECTO}”, ejecutado para {CLIENTE} conforme al "
    f"alcance y los términos acordados. A continuación se resumen los datos "
    f"generales del cierre.", after=12)

# Datos generales
add_kv_table([
    ("Proyecto", PROYECTO),
    ("Cliente", CLIENTE),
    ("Periodo de ejecución", PERIODO),
    ("Patrocinador / Contacto", PATROCINADOR),
    ("Número de contrato / OC", CONTRATO_OC),
])

# ═══════════════════════════════════════════════════════════════════
# OBJETIVOS ALCANZADOS
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Objetivos alcanzados")
add_paragraph("Durante la ejecución del proyecto se cumplieron los siguientes objetivos:", after=6)
for obj in OBJETIVOS:
    add_bullet(obj)

# ═══════════════════════════════════════════════════════════════════
# ENTREGABLES
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Entregables")
add_branded_table(
    ["Entregable", "Fecha de entrega", "Aceptado"],
    ENTREGABLES,
    col_widths=[Inches(3.6), Inches(1.7), Inches(1.2)],
)

# ═══════════════════════════════════════════════════════════════════
# PENDIENTES / OBSERVACIONES
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Asuntos pendientes / observaciones")
add_paragraph(PENDIENTES, after=10)

# ═══════════════════════════════════════════════════════════════════
# DECLARACIÓN DE CIERRE
# ═══════════════════════════════════════════════════════════════════
add_heading_bar("Declaración de cierre")
add_paragraph(
    "Con la firma de la presente, ambas partes reconocen que los entregables "
    "fueron recibidos a satisfacción y que el proyecto se da por concluido "
    "formalmente. Agradecemos la confianza depositada en Mobiik y reiteramos "
    "nuestra disposición para colaborar en iniciativas futuras.", after=18)

add_paragraph("Atentamente,", after=4)

# ═══════════════════════════════════════════════════════════════════
# FIRMAS
# ═══════════════════════════════════════════════════════════════════
add_signature_block(*FIRMA_MOBIIK)
add_signature_block(*FIRMA_CLIENTE)

# Pie de página
doc.add_paragraph()
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run("© 2026 Mobiik   |   coding for the future   |   Documento Confidencial")
rf.font.name = FONT; rf.font.size = Pt(9); rf.font.italic = True; rf.font.color.rgb = GREY


# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
import os
os.makedirs(os.path.dirname(OUT) or ".", exist_ok=True)
doc.save(OUT)
print(f"✓ Saved: {OUT}")
